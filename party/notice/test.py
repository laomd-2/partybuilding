from docx import Document
from docx.oxml.ns import qn


document = Document('材料21：接收预备党员备案表.docx')
table  = document.tables[0]
print(table.__dict__)
# document.styles['Normal'].font.name = '微软雅黑'
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# rows_num = 4
# cols_num = 12
# table = document.add_table(rows=rows_num, cols=cols_num, style = 'Table Grid')

# for r in range(rows_num):
    # for c in range(cols_num):
        # print(table.cell(r, c).text, end=' ')
    # print()
#         table.cell(r, c).text = "第{r}行{c}列".format(r = r+1, c = c+1)
# # 保存文档
# document.save('Python生成的文档.docx')