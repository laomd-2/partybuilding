from django.contrib import messages
from django.core import mail
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse, HttpResponseNotFound, HttpResponseRedirect
from django.shortcuts import render
from django.utils.encoding import escape_uri_path

from notice.util import _queryset
from info.models import get_branch_managers
from info.util import group_by_branch
from notice.adminx import *
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL


class PlanView(ListAdminView):
    model = Member
    table_mapping = dict([(model.__name__.lower(), model)
                          for model in [FirstTalk, Activist, KeyDevelop, LearningClass, PreMember, FullMember]])

    def get_context(self):
        context = super().get_context()
        context['can_send_email'] = is_school_admin(self.request.user)
        context['can_beian'] = is_admin(self.request.user)
        return context

    def get(self, request, *args, **kwargs):
        phase = request.GET.get('phase')
        if phase is None or phase not in self.table_mapping:
            return HttpResponseNotFound('未找到资源。')
        else:
            model = self.table_mapping[phase]
            query = self.result_list = _queryset(request, model)
            self.ordering_field_columns = model.fields
            context = self.get_context()
            # context['headers'] = get_headers(model.fields)
            # context['results'] = [[m[f] for f in model.fields] for m in query]
            return render(request, 'plan/%s.html' % phase, context)


class TableView(PlanView):

    def get(self, request, table):
        model = self.table_mapping.get(table)
        if model is not None:
            return self.export(request, model)
        return HttpResponseNotFound('未找到资源')

    @staticmethod
    def export(request, model):
        filename = model.export_filename()
        data = model.export(request)
        response = HttpResponse(data, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response['Content-Disposition'] = "attachment; filename*=utf-8''{}".format(escape_uri_path(filename))
        return response


class BeianView(TableView):

    @staticmethod
    def export(request, model):
        query = queryset(request, model, fields=model.beian_fields)
        filename = model.beian_template

        doc = Document(filename)
        table = doc.tables[0]

        try:
            UserStyle1 = doc.styles.add_style('UserStyle1', 1)
        except ValueError:
            UserStyle1 = doc.styles['UserStyle1']
        # 设置字体尺寸
        UserStyle1.font.size = Pt(12)
        # 设置中文字体
        UserStyle1.font.name = '仿宋_GB2312'
        UserStyle1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        for row in query:
            row['birth_date'] = '-'.join(str(row['birth_date']).split('-')[:-1])
            tb_row = table.add_row()
            values = [row[k] for k in model.beian_fields]
            model.complete_beian(values)
            for v, cell in zip(values, tb_row.cells):
                cell.text = wrap(v)
                for p in cell.paragraphs:
                    p.style = UserStyle1
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        data = to_bytes(doc)
        filename = filename.split('/')[-1]
        response = HttpResponse(data,
                                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response['Content-Disposition'] = "attachment; filename*=utf-8''{}".format(escape_uri_path(filename))
        return response


class EmailView(PlanView):

    def post(self, request, table):
        model = self.table_mapping.get(table)
        if model is not None:
            return self.send_email(request, model, model.verbose_name, model.verbose_name, model.phase)
        return HttpResponseNotFound('NOT FOUND')

    @staticmethod
    def send_email(request, model, manager_title, member_title, phase):
        if not is_school_admin(request.user):
            raise PermissionDenied
        groups = group_by_branch(queryset(request, model))
        branch_managers = get_branch_managers()
        fields = model.fields
        mails = []
        success = []
        for branch, appers in groups.items():
            appers = list(appers)
            if branch in branch_managers:
                email = make_email_to_managers(branch_managers[branch], manager_title, appers,
                                                    fields, phase)
                if email:
                    mails.append(email)
                    success.append(branch)
            # mails.extend(make_email_to_appliers(member_title, appers, fields))
        connection = mail.get_connection()  # Use default email connection
        connection.fail_silently = True
        cnt = connection.send_messages(mails)
        messages.success(request, '%s：向%s发送邮件，%d封发送成功！' % (manager_title, ','.join(success), cnt))
        return HttpResponseRedirect('/')
